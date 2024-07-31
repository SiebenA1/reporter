# -*- coding: utf-8 -*-
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from report_generator.compontent.settings import SETTINGS


def insert_page_break(doc: Document) -> None:
    """
    Insert a page break to the document
    """
    # Create a new paragraph for the page break
    page_break_paragraph = doc.add_paragraph()

    # Add the page break to the new paragraph
    run = page_break_paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    # Ensure the paragraph does not contain any text to prevent unnecessary blank pages
    page_break_paragraph.clear()


def string_to_rgb_color(rgb_string: str) -> RGBColor:
    """
    Convert a string to RGBColor

    Parameters
    ----------
    rgb_string : str

    Returns
    -------
    RGBColor
    """
    rgb_values = rgb_string.strip('()').split(',')

    r = int(rgb_values[0], 16)
    g = int(rgb_values[1], 16)
    b = int(rgb_values[2], 16)

    return RGBColor(r, g, b)


def set_global_formatting(doc: Document) -> None:
    """
    Set the global formatting of the document

    Parameters
    ----------

    doc : Document
        The document to set the formatting
    """
    # Get the settings for the formatting
    line_spacing = SETTINGS.get('line_spacing')
    top_margin = SETTINGS.get('top_margin')
    bottom_margin = SETTINGS.get('bottom_margin')
    left_margin = SETTINGS.get('left_margin')
    right_margin = SETTINGS.get('right_margin')
    # Set the margins
    section = doc.sections[0]
    section.top_margin = Inches(top_margin if top_margin is not None else 1.0)
    section.bottom_margin = Inches(bottom_margin if bottom_margin is not None else 1.0)
    section.left_margin = Inches(left_margin if left_margin is not None else 1.0)
    section.right_margin = Inches(right_margin if right_margin is not None else 1.0)
    # Set the line spacing
    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = Pt((line_spacing if line_spacing is not None else 1.0) * 12)


def add_page_number(paragraph: Paragraph) -> None:
    """
    Add page number to the paragraph

    Parameters
    ----------

    paragraph : Paragraph
        The paragraph to add page number
    """
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    t = OxmlElement('w:t')
    t.text = "1"
    run._r.append(t)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

    run.add_text(' / ')

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar4)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    run._r.append(instrText2)

    fldChar5 = OxmlElement('w:fldChar')
    fldChar5.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar5)

    t2 = OxmlElement('w:t')
    t2.text = "1"
    run._r.append(t2)

    fldChar6 = OxmlElement('w:fldChar')
    fldChar6.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar6)

    run.font.size = Pt(9)
