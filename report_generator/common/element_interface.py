# -*- coding: utf-8 -*-
import json
from abc import ABC
from pathlib import Path
from typing import Tuple

from document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, Inches, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from report_generator.compontent.global_setting_interface import add_page_number, string_to_rgb_color
from report_generator.compontent.settings import TEXT_FORMAT


class Element(ABC):
    """
    Base class for all elements in the document
    """

    def render(self, document: Document) -> None:
        pass


class TextFormat(Element):
    """
    Text format, including font name, font size, bold, italic, alignment
    """

    def __init__(self, FORMAT=TEXT_FORMAT['PARAGRAPH']):
        """
        Initialize the parameters of text format

        Parameters
        ----------
        FORMAT : dict
            The format of the text, including font name, font size, bold, italic, alignment, color, line
            spacing
        """
        self.font_name = FORMAT.get('font_name')
        self.font_size = FORMAT.get('font_size')
        self.bold = False if FORMAT.get('bold') == 'False' else True
        self.italic = False if FORMAT.get('italic') == 'False' else True
        self.color = string_to_rgb_color(FORMAT.get('color'))
        self.line_spacing = FORMAT.get('line_spacing') if FORMAT.get('line_spacing') else 1.5
        self.alignment = WD_ALIGN_PARAGRAPH.CENTER if FORMAT.get('alignment') == 'center' else WD_ALIGN_PARAGRAPH.LEFT

    def apply_format(self, run) -> None:
        """
        Apply the format to the run

        Parameters
        ----------
        run : docx.text.run.Run
        """
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run.font.color.rgb = self.color
        run.bold = self.bold
        run.italic = self.italic


class TitleTextFormat(TextFormat):
    """
    Title text format
    """

    def __init__(self, level: int):
        """
        Initialize the title text format with different levels

        Parameters
        ----------
        level : int
            The level of the title
        """
        if level == 1:
            super().__init__(FORMAT=TEXT_FORMAT['TITLE']["L1"])
        elif level == 2:
            super().__init__(FORMAT=TEXT_FORMAT['TITLE']["L2"])
        elif level == 3:
            super().__init__(FORMAT=TEXT_FORMAT['TITLE']["L3"])


class NormalTextFormat(TextFormat):
    """
    Normal text format
    """

    def __init__(self):
        """
        Initialize the normal text format
        """
        super().__init__(FORMAT=TEXT_FORMAT['PARAGRAPH'])


class PositiveStatusTextFormat(TextFormat):
    """
    Status text format for True status
    """

    def __init__(self):
        super().__init__(FORMAT=TEXT_FORMAT['POSITIVE_STATUS'])


class NegativeStatusTextFormat(TextFormat):
    """
    Status text format for False status
    """

    def __init__(self):
        super().__init__(FORMAT=TEXT_FORMAT['NEGATIVE_STATUS'])


class CaptionTextFormat(TextFormat):
    """
    Caption text format
    """

    def __init__(self):
        super().__init__(FORMAT=TEXT_FORMAT['CAPTION'])


class TableTextFormat(TextFormat):
    """
    Table text format
    """

    def __init__(self):
        super().__init__(FORMAT=TEXT_FORMAT['TABLE'])


class HeaderTextFormat(TextFormat):
    """
    Header text format
    """

    def __init__(self):
        super().__init__(FORMAT=TEXT_FORMAT['HEADER'])


class FooterTextFormat(TextFormat):
    """
    Footer text format
    """

    def __init__(self):
        super().__init__(FORMAT=TEXT_FORMAT['FOOTER'])


class Title(Element):
    """
    Title element
    """

    def __init__(self, text: str, level: int):
        """
        Initialize the title element

        Parameters
        ----------
        text : str
            The text of the title
        level : int
            The level of the title
        """
        self.text = text
        self.level = level
        if self.level == 1:
            self.text_format = TitleTextFormat(level=1)
        elif self.level == 2:
            self.text_format = TitleTextFormat(level=2)
        elif self.level == 3:
            self.text_format = TitleTextFormat(level=3)

    def render(self, document: Document) -> None:
        """
        Render the title element

        Parameters
        ----------

        document : docx.document.Document
        """
        document.add_heading(self.text, level=self.level)
        if self.text_format:
            title = document.paragraphs[-1]
            self.text_format.apply_format(title.runs[0])


class Paragraph(Element):
    """
    Paragraph element, including text and text format
    """

    def __init__(self, title: str, text: str, text_format: TextFormat):
        """
        Initialize the paragraph element

        Parameters
        ----------
        title : str
            The title of the paragraph
        text : str
            The text of the paragraph
        text_format : object
            The text format of the paragraph
        """
        self.title = title
        self.text = text
        self.text_format = text_format if text_format else NormalTextFormat()

    def render(self, document: Document) -> None:
        """
        Render the paragraph element

        Parameters
        ----------
        document : docx.document.Document
        """
        if self.title:
            document.add_heading(self.title, level=2)
        else:
            document.add_paragraph()
        p = document.add_paragraph(self.text)
        if self.text_format:
            run = p.runs[0]
            self.text_format.apply_format(run)
            p.alignment = self.text_format.alignment


class Image(Element):
    """
    Image element, including an image path
    """

    def __init__(self, case_name: str, image_path: Path, width=None, height=None):
        self.case_name = case_name
        self.path = image_path  # path of the image
        self.width = width  # width of the image, in inches
        self.height = height  # height of the image, in inches

    def render(self, document: Document) -> None:
        """
        Render the image element

        Parameters
        ----------
        document : docx.document.Document
        """
        section = document.sections[0]
        page_width = section.page_width - section.left_margin - section.right_margin
        img_width = Inches(self.width) if self.width else page_width
        img_height = Inches(self.height) if self.height else None

        with open(self.path, 'r') as f:
            image_data = json.load(f)

        for case_data in image_data:
            if self.case_name in case_data:
                files = case_data[self.case_name]
                for file_name, image_paths in files.items():
                    title_text = f"{self.case_name} - {file_name}"
                    title = document.add_heading(title_text, level=2)
                    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for image_path in image_paths:
                        document.add_picture(image_path, width=img_width, height=img_height)
                        last_paragraph = document.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


class Table(Element):
    """
    Table element, including data and title
    """

    def __init__(self, data: list, title=None):
        """
        Initialize the table element

        Parameters
        ----------
        data : list
            The path to the data file
        """
        self.title = title
        self.data = data
        self.line_spacing = TableTextFormat().line_spacing if TableTextFormat().line_spacing else 1.0
        self.text_format = TableTextFormat()

    def render(self, document: Document) -> None:
        if self.title:
            document.add_heading(self.title, level=2)

        rows = len(self.data)
        cols = len(self.data[0]) if rows > 0 else 0
        table = document.add_table(rows=rows, cols=cols)

        # increase the weight of the first columns
        for cell in table.columns[0].cells:
            cell.width = Inches(8.0)

        for i, row_data in enumerate(self.data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                self._set_cell_border(cell)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph_format.line_spacing = self.line_spacing
                run = paragraph.add_run(str(cell_data))
                self.text_format.apply_format(run)
                # the first row aligns right, the other rows align left
                if j == 0:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    @staticmethod
    def _set_cell_border(cell):
        """
        Set the border of the cell

        Parameters
        ----------
        cell : docx.table._Cell
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        for line in ['top', 'start', 'bottom', 'end']:
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            line_border = OxmlElement(f'w:{line}')
            line_border.set(qn('w:val'), 'single')
            line_border.set(qn('w:sz'), '4')
            line_border.set(qn('w:space'), '0')
            line_border.set(qn('w:color'), '000000')
            tcBorders.append(line_border)


class Tables(Element):
    def __init__(self, condition_result: dict):
        """
        Initialize the tables element

        Parameters
        ----------
        condition_result : dict
            The condition results
        """
        self.condition_result = condition_result

    def render(self, document: Document) -> None:
        """
        Render the tables element

        Parameters
        ----------
        document : docx.document.Document
        """
        for file_key, condition_list in self.condition_result.items():
            title = f"{file_key}: {'Passed' if all(result for _, result in condition_list) else 'Failed'}"
            condition_result_list = self._format_condition_result(condition_list)
            # Add elements to the document
            document.add_heading(title, level=2)
            table = Table(data=condition_result_list)
            table.render(document)

    @staticmethod
    def _format_condition_result(condition_result: list[Tuple[list[str], bool]]) -> list[list[str]]:
        """
        Format the condition results to a list of list to adapt to the table format

        Parameters
        ----------
        condition_result : List[Tuple[List[str], bool]]
            Condition results

        Returns
        -------
        List[List[str]]
            Formatted condition results
        """
        formatted_result = [["Expected Result", "Result"]]
        for condition, result in condition_result:
            conditions = ', '.join(condition)
            result_text = "Passed" if result else "Failed"
            formatted_result.append([conditions, result_text])
        return formatted_result


class SetupBuilder(Element):
    """
    Base class for global setup builder
    """

    def header_render(self, document: Document):
        pass

    def footer_render(self, document: Document):
        pass


class GlobalSetupBuilder(SetupBuilder):
    """
    Global setup builder, including headers and footers' setup
    """

    def __init__(self, doc: Document, left_header_text: str, footer_text: str, middle_footer_text: str,
                 image_path: Path):
        """
        Initialize the global setup builder

        Parameters
        ----------

        doc : docx.document.Document
            The document object
        left_header_text : str
            The text of the left header
        footer_text : str
            The text of the footer
        middle_footer_text : str
            The text of the middle footer
        image_path : Path
            The path to the image
        """
        self.document = doc
        self.image = image_path
        self.left_header_text = left_header_text
        self.footer_text = footer_text
        self.middle_footer_text = middle_footer_text
        self.header_text_format = HeaderTextFormat()
        self.footer_text_format = FooterTextFormat()

    def header_render(self, document: Document) -> 'GlobalSetupBuilder':
        """
        Render the header by rending a table with two cells, which are the left header and the right header,
        the left header is the text, and the right header is the image as the logo.

        Parameters
        ----------
        document : docx.document.Document

        Returns
        -------
        object
        """
        for section in self.document.sections:
            # header setting for each section
            header = section.header
            header.is_linked_to_previous = False
            # calculate the width of the page, left margin, and right margin
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            # calculate the width of the table
            table_width = page_width - left_margin - right_margin
            # add the table to the header
            table = header.add_table(rows=1, cols=2, width=Emu(table_width))
            table.allow_autofit = False
            # set the width of the columns
            table.columns[0].width = Inches(6.5)
            table.columns[1].width = Inches(1.0)
            # add the text to the left cell and render the text format
            left_cell = table.cell(0, 0)
            left_cell._element.clear_content()
            left_header_paragraph = left_cell.add_paragraph(self.left_header_text)
            if self.header_text_format:
                self.header_text_format.apply_format(left_header_paragraph.runs[0])
            left_header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            # add the image to the right cell
            if self.image:
                right_cell = table.cell(0, 1)
                right_cell._element.clear_content()
                header_paragraph = right_cell.add_paragraph()
                run = header_paragraph.add_run()
                run.add_picture(self.image, width=Inches(1.0))
                header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

        return self

    def footer_render(self, document: Document) -> 'GlobalSetupBuilder':
        """
        Render the footer by rendering a table with three cells, which are the left footer, the middle footer, and the right footer,
        the left footer is the text, the middle footer is the text, and the right footer is the page number.

        Parameters
        ----------
        document : docx.document.Document

        Returns
        -------
        object
        """
        for section in self.document.sections:
            # footer setting for each section
            footer = section.footer
            footer.is_linked_to_previous = False
            # calculate the width of the page as the width of the table
            table = footer.add_table(rows=1, cols=3, width=Emu(section.page_width))
            table.allow_autofit = False
            # set the width of the columns
            table.columns[0].width = Inches(5.5)
            table.columns[1].width = Inches(2.0)
            table.columns[2].width = Inches(1.0)
            # add the text to the left cell and render the text format
            left_cell = table.cell(0, 0)
            self._clear_cell_content(left_cell)
            footer_paragraph = left_cell.add_paragraph(self.footer_text)
            if self.footer_text_format:
                self.footer_text_format.apply_format(footer_paragraph.runs[0])
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # add the text to the middle cell and render the text format
            middle_cell = table.cell(0, 1)
            self._clear_cell_content(middle_cell)
            middle_paragraph = middle_cell.add_paragraph(self.middle_footer_text)
            if self.footer_text_format:
                self.footer_text_format.apply_format(middle_paragraph.runs[0])
            middle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # add the page number to the right cell
            right_cell = table.cell(0, 2)
            self._clear_cell_content(right_cell)
            page_number_paragraph = right_cell.add_paragraph()
            page_number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            add_page_number(page_number_paragraph)

        return self

    @staticmethod
    def _clear_cell_content(cell) -> None:
        """
        Clear the content of the cell
        """
        cell._element.clear_content()
